"""
========================================
Word 文档解析器
========================================

📚 模块说明：
- 解析Word文档（.docx, .doc）
- 提取文本、表格、图片等元素
- 保留文档结构信息

🎯 核心功能：
1. 文本提取（保留段落结构）
2. 表格识别和提取
3. 样式信息提取
4. 元数据提取

========================================
"""
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.logger import logger, log_execution
from core.config import settings


class WordParser:
    """
    Word文档解析器

    🎯 职责：
    - 提取Word文档文本
    - 提取表格数据
    - 提取文档元数据
    - 保留文档结构
    """

    def __init__(self):
        """初始化Word解析器"""
        self.supported_extensions = ['.docx', '.doc']

    @log_execution("解析Word文档")
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析Word文件

        参数：
            file_path: Word文件路径

        返回：
            Dict: 解析结果
            {
                "text": str,              # 全文本
                "paragraphs": List[Dict], # 段落列表
                "tables": List[Dict],     # 表格数据
                "metadata": Dict,         # 文档元数据
                "has_tables": bool,       # 是否包含表格
                "paragraph_count": int,   # 段落数
                "word_count": int,        # 字数
            }
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            logger.info(f"开始解析Word: {file_path}")

            # 打开Word文档
            doc = Document(file_path)

            # 提取元数据
            metadata = self._extract_metadata(doc)

            # 提取段落
            paragraphs_data = self._extract_paragraphs(doc)

            # 提取表格
            tables = self._extract_tables(doc)

            # 合并所有文本
            all_text = "\n\n".join([p["text"] for p in paragraphs_data if p["text"]])

            # 统计字数
            word_count = len(all_text.replace(" ", "").replace("\n", ""))

            result = {
                "text": all_text,
                "paragraphs": paragraphs_data,
                "tables": tables,
                "metadata": metadata,
                "has_tables": len(tables) > 0,
                "paragraph_count": len(paragraphs_data),
                "word_count": word_count,
                "file_path": file_path,
                "file_name": os.path.basename(file_path)
            }

            logger.info(
                f"Word解析完成: {file_path} | "
                f"段落: {len(paragraphs_data)} | "
                f"表格: {len(tables)} | "
                f"字数: {word_count}"
            )

            return result

        except Exception as e:
            logger.error(f"Word解析失败: {file_path} | 错误: {str(e)}")
            raise

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """
        提取Word文档元数据

        参数：
            doc: Document对象

        返回：
            Dict: 元数据信息
        """
        try:
            # Word核心属性
            core_props = doc.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
            }

            return metadata

        except Exception as e:
            logger.warning(f"提取Word元数据失败: {str(e)}")
            return {}

    @staticmethod
    def _is_heading_style(style_name: str) -> bool:
        """判断样式是否为标题（兼容英文 Heading 与中文 标题）"""
        if not style_name:
            return False
        return style_name.startswith("Heading") or style_name.startswith("标题")

    @staticmethod
    def _heading_level(style_name: str) -> int:
        """从样式名提取标题级别，无数字时返回 0"""
        digits = "".join(ch for ch in style_name if ch.isdigit())
        try:
            return int(digits) if digits else 0
        except (ValueError, TypeError):
            return 0

    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取段落信息

        参数：
            doc: Document对象

        返回：
            List[Dict]: 段落数据列表

        💡 段落格式：
        {
            "index": int,         # 段落索引
            "text": str,          # 段落文本
            "style": str,         # 样式名称
            "is_heading": bool,   # 是否为标题
            "level": int,         # 标题级别（如果是标题）
        }
        """
        paragraphs_data = []

        try:
            for index, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                # 跳过空段落
                if not text:
                    continue

                # 获取样式信息
                style_name = paragraph.style.name if paragraph.style else "Normal"

                # 判断是否为标题（兼容中文样式名 "标题 1"）
                is_heading = self._is_heading_style(style_name)

                # 提取标题级别
                level = self._heading_level(style_name) if is_heading else 0

                para_data = {
                    "index": index,
                    "text": text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "level": level
                }

                paragraphs_data.append(para_data)

            return paragraphs_data

        except Exception as e:
            logger.error(f"提取Word段落失败: {str(e)}")
            raise

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取Word文档中的表格

        参数：
            doc: Document对象

        返回：
            List[Dict]: 表格数据列表

        💡 表格格式：
        {
            "table_index": int,   # 表格索引
            "rows": int,          # 行数
            "cols": int,          # 列数
            "data": List[List],   # 表格数据
            "text": str           # 表格转换为文本
        }
        """
        tables = []

        try:
            for table_index, table in enumerate(doc.tables):
                # 提取表格数据
                table_data = []

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)

                # 转换为文本
                table_text = self._table_to_text(table_data)

                table_info = {
                    "table_index": table_index,
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0,
                    "data": table_data,
                    "text": table_text
                }

                tables.append(table_info)

            logger.info(f"提取到 {len(tables)} 个表格")
            return tables

        except Exception as e:
            logger.warning(f"提取Word表格失败: {str(e)}")
            return []

    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """
        将表格数据转换为文本格式

        参数：
            table_data: 表格数据（二维列表）

        返回：
            str: 格式化的表格文本
        """
        if not table_data:
            return ""

        try:
            lines = []
            for row in table_data:
                # 用 | 分隔单元格
                line = " | ".join(row)
                lines.append(line)

            return "\n".join(lines)

        except Exception as e:
            logger.warning(f"表格转文本失败: {str(e)}")
            return ""

    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        提取文档中的所有标题

        参数：
            file_path: Word文件路径

        返回：
            List[Dict]: 标题列表

        💡 用途：
        - 生成文档目录
        - 理解文档结构
        - 辅助文本分块
        """
        try:
            doc = Document(file_path)
            headings = []

            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name if paragraph.style else "Normal"
                if self._is_heading_style(style_name):
                    headings.append({
                        "text": paragraph.text.strip(),
                        "level": self._heading_level(style_name),
                        "style": style_name
                    })

            return headings

        except Exception as e:
            logger.error(f"提取标题失败: {str(e)}")
            return []

    def get_text_with_structure(self, file_path: str) -> str:
        """
        获取保留结构的文本

        参数：
            file_path: Word文件路径

        返回：
            str: 带结构标记的文本

        💡 格式示例：
        # 标题1
        正文段落...

        ## 标题2
        正文段落...

        [表格]
        | 列1 | 列2 |
        | 数据1 | 数据2 |
        """
        try:
            result = self.parse(file_path)

            structured_text = []

            # 添加段落（标题用#标记）
            for para in result['paragraphs']:
                if para['is_heading']:
                    # 标题前加#号
                    prefix = "#" * para['level']
                    structured_text.append(f"{prefix} {para['text']}")
                else:
                    structured_text.append(para['text'])

            # 添加表格
            for table in result['tables']:
                structured_text.append("\n[表格]")
                structured_text.append(table['text'])

            return "\n\n".join(structured_text)

        except Exception as e:
            logger.error(f"获取结构化文本失败: {str(e)}")
            raise


# =========================================
# 💡 使用示例
# =========================================
"""
# 1. 基础使用
from services.document.word_parser import WordParser

parser = WordParser()

# 解析Word文档
result = parser.parse("data/raw_docs/项目总结.docx")

print(f"文档标题: {result['metadata']['title']}")
print(f"作者: {result['metadata']['author']}")
print(f"段落数: {result['paragraph_count']}")
print(f"字数: {result['word_count']}")
print(f"表格数: {len(result['tables'])}")

# 访问全文本
full_text = result['text']
print(full_text[:500])

# 访问段落
for para in result['paragraphs'][:5]:
    if para['is_heading']:
        print(f"[标题{para['level']}] {para['text']}")
    else:
        print(para['text'][:100])

# 访问表格
for table in result['tables']:
    print(f"\n表格 {table['table_index']}:")
    print(table['text'])


# 2. 提取标题（生成目录）
headings = parser.extract_headings("document.docx")
for heading in headings:
    indent = "  " * (heading['level'] - 1)
    print(f"{indent}- {heading['text']}")


# 3. 获取结构化文本
structured_text = parser.get_text_with_structure("document.docx")
print(structured_text)
"""